"""Document parsing service for extracting text from various file formats."""

import csv
import io
from pathlib import Path
from typing import Optional


class DocumentParser:
    """Parse various document formats and extract text content."""

    MAX_TEXT_LENGTH = 100000  # Limit text to ~100k chars for API limits

    async def parse_file(self, file_path: Path, mime_type: str) -> dict:
        """
        Parse a file and extract its text content.

        Returns:
            dict with 'text', 'pages' (if applicable), 'error' (if failed)
        """
        try:
            if mime_type == "application/pdf":
                return await self._parse_pdf(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._parse_docx(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                return await self._parse_xlsx(file_path)
            elif mime_type == "application/vnd.ms-excel":
                return await self._parse_xls(file_path)
            elif mime_type == "text/csv":
                return await self._parse_csv(file_path)
            elif mime_type in ["text/plain", "text/markdown"]:
                return await self._parse_text(file_path)
            else:
                return {"text": "", "error": f"Unsupported file type: {mime_type}"}
        except Exception as e:
            return {"text": "", "error": str(e)}

    async def _parse_pdf(self, file_path: Path) -> dict:
        """Extract text from PDF files."""
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages = []
        full_text = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page": i + 1, "text": text})
            full_text.append(f"--- Page {i + 1} ---\n{text}")

        combined_text = "\n\n".join(full_text)
        return {
            "text": combined_text[:self.MAX_TEXT_LENGTH],
            "pages": len(pages),
            "truncated": len(combined_text) > self.MAX_TEXT_LENGTH,
        }

    async def _parse_docx(self, file_path: Path) -> dict:
        """Extract text from DOCX files."""
        from docx import Document

        doc = Document(str(file_path))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        combined_text = "\n\n".join(paragraphs)
        return {
            "text": combined_text[:self.MAX_TEXT_LENGTH],
            "paragraphs": len(paragraphs),
            "truncated": len(combined_text) > self.MAX_TEXT_LENGTH,
        }

    async def _parse_xlsx(self, file_path: Path) -> dict:
        """Extract text from XLSX files."""
        from openpyxl import load_workbook

        wb = load_workbook(str(file_path), data_only=True)
        sheets_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []

            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to strings
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    rows.append(" | ".join(row_values))

            if rows:
                sheets_text.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))

        combined_text = "\n\n".join(sheets_text)
        return {
            "text": combined_text[:self.MAX_TEXT_LENGTH],
            "sheets": len(wb.sheetnames),
            "truncated": len(combined_text) > self.MAX_TEXT_LENGTH,
        }

    async def _parse_xls(self, file_path: Path) -> dict:
        """Extract text from XLS files (older Excel format)."""
        import xlrd

        wb = xlrd.open_workbook(str(file_path))
        sheets_text = []

        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            rows = []

            for row_idx in range(sheet.nrows):
                row_values = [
                    str(cell.value) for cell in sheet.row(row_idx)
                    if cell.value not in (None, "")
                ]
                if row_values:
                    rows.append(" | ".join(row_values))

            if rows:
                sheets_text.append(f"=== Sheet: {sheet.name} ===\n" + "\n".join(rows))

        combined_text = "\n\n".join(sheets_text)
        return {
            "text": combined_text[:self.MAX_TEXT_LENGTH],
            "sheets": wb.nsheets,
            "truncated": len(combined_text) > self.MAX_TEXT_LENGTH,
        }

    async def _parse_csv(self, file_path: Path) -> dict:
        """Extract text from CSV files."""
        rows = []

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                row_values = [cell for cell in row if cell.strip()]
                if row_values:
                    rows.append(" | ".join(row_values))

        combined_text = "\n".join(rows)
        return {
            "text": combined_text[:self.MAX_TEXT_LENGTH],
            "rows": len(rows),
            "truncated": len(combined_text) > self.MAX_TEXT_LENGTH,
        }

    async def _parse_text(self, file_path: Path) -> dict:
        """Extract text from plain text/markdown files."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        return {
            "text": text[:self.MAX_TEXT_LENGTH],
            "truncated": len(text) > self.MAX_TEXT_LENGTH,
        }


# Singleton instance
document_parser = DocumentParser()
